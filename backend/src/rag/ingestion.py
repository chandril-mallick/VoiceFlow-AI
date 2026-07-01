"""
VoiceFlow AI — RAG Document Ingestion
Process and index uploaded documents (PDF, DOCX, CSV, TXT, MD, URL) into Qdrant.
"""

import logging
from pathlib import Path
from typing import Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# Supported file types and their loaders
SUPPORTED_TYPES = {"pdf", "docx", "csv", "txt", "md", "url"}


def load_document(file_path: str, file_type: str) -> list[Document]:
    """
    Load a document from file path based on its type.

    Returns:
        List of LangChain Document objects.
    """
    path = Path(file_path)
    documents = []

    try:
        if file_type == "pdf":
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    documents.append(Document(
                        page_content=text,
                        metadata={"source": path.name, "page": i + 1, "type": "pdf"},
                    ))

        elif file_type == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(str(path))
            text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={"source": path.name, "type": "docx"},
                ))

        elif file_type == "csv":
            import csv
            with open(path, "r", encoding="utf-8") as f:
                reader_csv = csv.DictReader(f)
                for i, row in enumerate(reader_csv):
                    row_text = " | ".join(f"{k}: {v}" for k, v in row.items() if v)
                    if row_text:
                        documents.append(Document(
                            page_content=row_text,
                            metadata={"source": path.name, "row": i + 1, "type": "csv"},
                        ))

        elif file_type in ("txt", "md"):
            text = path.read_text(encoding="utf-8")
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"source": path.name, "type": file_type},
                ))

        elif file_type == "url":
            from bs4 import BeautifulSoup
            import httpx
            url = path.read_text().strip() if path.exists() else str(path)
            response = httpx.get(url, timeout=30, follow_redirects=True)
            soup = BeautifulSoup(response.text, "html.parser")
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer"]):
                script.decompose()
            text = soup.get_text(separator="\n", strip=True)
            if text:
                documents.append(Document(
                    page_content=text,
                    metadata={"source": url, "type": "url"},
                ))

    except Exception as e:
        logger.error("Failed to load %s (%s): %s", file_path, file_type, e)
        raise

    logger.info("Loaded %d document sections from %s", len(documents), file_path)
    return documents


def split_documents(
    documents: list[Document],
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
) -> list[Document]:
    """Split documents into smaller chunks for embedding."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", ", ", " ", ""],
    )

    chunks = text_splitter.split_documents(documents)
    logger.info("Split into %d chunks (size=%d, overlap=%d)", len(chunks), chunk_size, chunk_overlap)
    return chunks


async def process_document(
    file_path: str,
    file_type: str,
    tenant_id: str,
    collection_name: Optional[str] = None,
) -> dict:
    """
    Full document processing pipeline:
    1. Load document
    2. Split into chunks
    3. Generate embeddings
    4. Store in Qdrant

    Returns:
        dict with chunk_count, collection_name, status
    """
    from src.rag.embeddings import get_embedding_function
    from src.rag.retriever import get_qdrant_store

    # Default collection name per tenant
    if not collection_name:
        collection_name = f"tenant_{tenant_id}"

    # 1. Load
    documents = load_document(file_path, file_type)
    if not documents:
        return {"chunk_count": 0, "status": "empty", "collection_name": collection_name}

    # 2. Split
    chunks = split_documents(documents)

    # Add tenant metadata to all chunks
    for chunk in chunks:
        chunk.metadata["tenant_id"] = tenant_id

    # 3. Embed and store
    embeddings = get_embedding_function()
    store = get_qdrant_store(collection_name, embeddings)

    await store.aadd_documents(chunks)

    logger.info(
        "✅ Indexed %d chunks for tenant %s in collection %s",
        len(chunks), tenant_id, collection_name,
    )

    return {
        "chunk_count": len(chunks),
        "collection_name": collection_name,
        "status": "completed",
    }
